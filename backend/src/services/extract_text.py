import sys
import os

def extract(file_path):
    if not os.path.exists(file_path):
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    pass
            for page in reader.pages:
                pt = page.extract_text()
                if pt:
                    text += pt + "\n"
        except Exception as e:
            pass
            
    elif ext in [".docx", ".doc"]:
        try:
            import docx
            doc = docx.Document(file_path)
            for p in doc.paragraphs:
                if p.text.strip():
                    text += p.text.strip() + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_t = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_t:
                        text += " | ".join(row_t) + "\n"
        except Exception as e:
            pass
            
    return text.strip()

if __name__ == "__main__":
    if len(sys.argv) > 1:
        extracted = extract(sys.argv[1])
        # Write to stdout using utf-8 encoding
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stdout.write(extracted)
